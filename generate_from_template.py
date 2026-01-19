"""
Document Generator - Multi-Source JSON Template System

This generator creates Word documents using:
1. A template structure (static text with placeholder references)
2. Multiple JSON data sources (CFR, CYC, Illustration, Ceding Info, User Input)

The approach separates:
- Template content (what stays the same)
- Dynamic data by source (what changes per client/document)
"""

import json
import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import existing style helpers
from style_helpers import (
    apply_run_style, apply_paragraph_style, set_cell_background,
    set_vertical_text, apply_section_settings, set_row_height,
    DEFAULT_STYLES, hex_to_rgb
)


class DataSourceManager:
    """
    Manages multiple JSON data sources and provides unified access to placeholders.
    
    Sources:
    - cfr: Client Financial Record (GREEN)
    - cyc: Calculations (YELLOW) 
    - illustration: Illustration data (PINK)
    - ceding_info: Ceding Info Check (RED)
    - user_input: User Input (BLUE)
    """
    
    def __init__(self, data_dir="data"):
        self.data_dir = data_dir
        self.sources = {}
        self._load_all_sources()
    
    def _load_all_sources(self):
        """Load all data source JSON files."""
        source_files = {
            "cfr": "cfr_data.json",
            "cyc": "cyc_data.json",
            "illustration": "illustration_data.json",
            "ceding_info": "ceding_info.json",
            "user_input": "user_input.json"
        }
        
        for source_name, filename in source_files.items():
            filepath = os.path.join(self.data_dir, filename)
            if os.path.exists(filepath):
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.sources[source_name] = json.load(f)
                print(f"  ✓ Loaded {source_name} from {filename}")
            else:
                print(f"  ⚠ Warning: {filename} not found")
                self.sources[source_name] = {}
    
    def get(self, path, default=""):
        """
        Get a value from any source using dot notation.
        
        Examples:
            get("cfr.recipient.title_and_name")
            get("user_input.letter_details.date")
        """
        parts = path.split(".")
        if len(parts) < 2:
            return default
        
        source_name = parts[0]
        if source_name not in self.sources:
            print(f"  ⚠ Unknown source: {source_name}")
            return default
        
        value = self.sources[source_name]
        for part in parts[1:]:
            if isinstance(value, dict) and part in value:
                value = value[part]
            elif isinstance(value, list):
                try:
                    value = value[int(part)]
                except (ValueError, IndexError):
                    return default
            else:
                return default
        
        return value if value is not None else default
    
    def resolve_placeholders(self, text):
        """
        Replace all {source.path} placeholders in text with actual values.
        """
        if not isinstance(text, str):
            return text
        
        pattern = r'\{([a-z_]+\.[a-zA-Z0-9_.]+)\}'
        
        def replace_match(match):
            path = match.group(1)
            value = self.get(path)
            return str(value) if value else match.group(0)
        
        return re.sub(pattern, replace_match, text)


class TemplateDocumentGenerator:
    """
    Document generator that uses a template structure and multiple data sources.
    """
    
    def __init__(self, template_path="data/template_structure.json", data_dir="data"):
        print("\n[INIT] Initializing Template Document Generator...")
        
        # Load template structure
        with open(template_path, 'r', encoding='utf-8') as f:
            self.template = json.load(f)
        print(f"  ✓ Loaded template structure")
        
        # Load data sources
        self.data = DataSourceManager(data_dir)
        
        # Load styles
        self.styles = self.template.get("styles", DEFAULT_STYLES)
    
    def generate(self, output_path):
        """Generate the complete document."""
        print(f"\n[GENERATE] Creating document: {output_path}")
        
        # Create new document
        doc = Document()
        
        # Setup document styles and page settings
        self._setup_document(doc)
        
        # Generate letter section
        self._generate_letter(doc)
        
        # Generate Part 1
        self._generate_part1(doc)
        
        # Generate Part 2
        self._generate_part2(doc)
        
        # Generate Part 3
        self._generate_part3(doc)
        
        # Generate Part 4
        self._generate_part4(doc)
        
        # Generate Appendix I (Further details)
        self._generate_appendix_i(doc)
        
        # Generate Appendix II (Product Comparison)
        self._generate_appendix(doc)
        
        # Add page numbers
        self._add_page_numbers(doc)
        
        # Save document
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        print(f"\n[OK] Document saved: {output_path}")
    
    def _setup_document(self, doc):
        """Setup document-level styles and settings."""
        # Set page size to A4
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        
        # Get letter margins from template
        letter_margins = self.template.get("pages", {}).get("letter", {}).get("margins", {})
        section.top_margin = Inches(letter_margins.get("top", 1))
        section.bottom_margin = Inches(letter_margins.get("bottom", 1))
        section.left_margin = Inches(letter_margins.get("left", 1))
        section.right_margin = Inches(letter_margins.get("right", 1))
        
        # Set default Normal style
        style = doc.styles['Normal']
        doc_style = self.styles.get("document", {})
        style.font.name = doc_style.get("font_family", "Poppins")
        style.font.size = Pt(doc_style.get("font_size", 10))
        style.paragraph_format.line_spacing = doc_style.get("line_spacing", 1.0)
        style.paragraph_format.space_after = Pt(0)
    
    def _generate_letter(self, doc):
        """Generate the letter section (page 1-2)."""
        print("  Generating letter section...")
        
        letter_template = self.template.get("template", {})
        
        # Confidential header
        p = doc.add_paragraph(letter_template.get("letter", {}).get("confidential", "Private and Confidential"))
        self._apply_style(p, {"bold": True})
        
        # Recipient address (from CFR)
        recipient = self.data.sources.get("cfr", {}).get("recipient", {})
        self._add_paragraph(doc, recipient.get("title_and_name", ""))
        for key in ["address_line_1", "address_line_2", "address_line_3", "address_line_4", "address_line_5"]:
            if recipient.get(key):
                self._add_paragraph(doc, recipient.get(key))
        
        doc.add_paragraph()  # Empty line
        
        # Date (from User Input)
        date_text = self.data.get("user_input.letter_details.date")
        p = doc.add_paragraph(date_text)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Salutation
        client_name = self.data.get("user_input.letter_details.client_first_name")
        self._add_paragraph(doc, f"Dear {client_name},")
        
        doc.add_paragraph()
        
        # Subject
        p = doc.add_paragraph("RE: Retirement Planning")
        self._apply_style(p, {"bold": True})
        
        doc.add_paragraph()
        
        # Intro paragraphs (with placeholder resolution)
        intro_paragraphs = letter_template.get("intro_paragraphs", [])
        for para_text in intro_paragraphs:
            resolved = self.data.resolve_placeholders(para_text)
            self._add_paragraph(doc, resolved)
            doc.add_paragraph()
        
        # Parts list
        parts_intro = letter_template.get("parts_intro", "")
        self._add_paragraph(doc, parts_intro)
        
        parts_bullets = letter_template.get("parts_bullets", [])
        for bullet in parts_bullets:
            self._add_bullet(doc, bullet)
        
        doc.add_paragraph()
        
        # Appendix intro
        appendix_intro = letter_template.get("appendix_intro", "")
        self._add_paragraph(doc, appendix_intro)
        
        doc.add_paragraph()
        
        # Other Documentation section
        other_docs = letter_template.get("other_documentation", {})
        self._add_heading2(doc, other_docs.get("heading", ""))
        self._add_paragraph(doc, other_docs.get("intro", ""))
        for bullet in other_docs.get("bullets", []):
            self._add_bullet(doc, bullet)
        doc.add_paragraph()
        self._add_paragraph(doc, other_docs.get("closing", ""))
        
        doc.add_paragraph()
        
        # Ongoing Advice section
        ongoing = letter_template.get("ongoing_advice", {})
        self._add_heading2(doc, ongoing.get("heading", ""))
        for para_text in ongoing.get("paragraphs", []):
            resolved = self.data.resolve_placeholders(para_text)
            self._add_paragraph(doc, resolved)
            doc.add_paragraph()
        
        # Next Steps section
        next_steps = letter_template.get("next_steps", {})
        self._add_heading2(doc, next_steps.get("heading", ""))
        for para_text in next_steps.get("paragraphs", []):
            resolved = self.data.resolve_placeholders(para_text)
            self._add_paragraph(doc, resolved)
            doc.add_paragraph()
        
        # Closing
        self._add_paragraph(doc, letter_template.get("letter", {}).get("valediction", "Yours sincerely,"))
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Adviser signature
        adviser_name = self.data.get("user_input.closing.adviser_name")
        adviser_title = self.data.get("user_input.closing.adviser_title")
        p = doc.add_paragraph()
        run = p.add_run(adviser_name)
        run.font.bold = True
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        
        self._add_paragraph(doc, adviser_title)
        self._add_paragraph(doc, letter_template.get("letter", {}).get("closing_company", ""))
    
    def _generate_part1(self, doc):
        """Generate Part 1 - Your Objectives, Needs and Circumstances."""
        print("  Generating Part 1...")
        
        part1 = self.template.get("part1", {})
        
        # Page break and Heading 1
        doc.add_page_break()
        self._add_heading1(doc, part1.get("heading1", "Part 1 - Your Objectives, Needs and Circumstances"))
        
        doc.add_paragraph()
        self._add_heading2(doc, part1.get("heading2", "Your Immediate Objectives"))
        
        # Process content items
        for item in part1.get("content", []):
            self._process_content_item(doc, item)
        
        # Scottish Widows Group Personal Pension details section
        self._generate_scottish_widows_details(doc)
        
        # LSA and LSDBA section
        self._generate_lsa_lsdba_section(doc)
    
    def _generate_scottish_widows_details(self, doc):
        """Generate the Scottish Widows plan details section."""
        plan_name = self.data.get("cfr.pension_arrangements.scottish_widows.provider_plan")
        
        doc.add_paragraph()
        self._add_heading2(doc, plan_name)
        
        # Get ceding info details
        ceding = self.data.sources.get("ceding_info", {}).get("scottish_widows_details", {})
        
        # Basic plan info bullets
        self._add_bullet(doc, ceding.get("plan_commenced", ""))
        self._add_bullet(doc, ceding.get("contribution_status", ""))
        self._add_bullet(doc, ceding.get("last_contribution", ""))
        self._add_bullet(doc, ceding.get("transfers_in", ""))
        self._add_bullet(doc, ceding.get("bulk_transfer", ""))
        self._add_bullet(doc, ceding.get("protected_tfc", ""))
        
        # Investments
        self._add_bullet(doc, "Your plan is invested in as follows:")
        investments = ceding.get("investments", [])
        for inv in investments:
            # Nested bullet with "o"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.75)
            run = p.add_run(f"o   {inv.get('fund_name', '')}                 {inv.get('value', '')}")
            run.font.name = "Poppins"
            run.font.size = Pt(10)
        
        # Charges
        self._add_bullet(doc, "The only charges applicable to the plan are:")
        charges = ceding.get("charges", {})
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        run = p.add_run(f"o   Annual management charge                              {charges.get('annual_management_charge', '')}")
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        run = p.add_run(f"o   Fund charge                                           {charges.get('fund_charge', '')}")
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        
        # Additional info bullets
        additional_info = ceding.get("additional_info", [])
        for info in additional_info:
            self._add_bullet(doc, info)
    
    def _generate_lsa_lsdba_section(self, doc):
        """Generate the LSA and LSDBA section."""
        doc.add_paragraph()
        self._add_heading2(doc, "Lump Sum Allowance (LSA) and Lump Sum and Death Benefit Allowance (LSDBA)")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "The Lump Sum Allowance (LSA) is the maximum amount of tax-free lump sum which could be paid from your pensions during your lifetime.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "The Lump Sum and Death Benefit Allowance (LSDBA) is the maximum tax-free lump sum which can be paid from your pensions during your lifetime (but this is usually first limited by your LSA) and on your death if you die before age 75.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "When drawing pension benefits during your lifetime, the tax-free lump sum available is dependent on the value of your pension, as well as your remaining LSA and LSDBA at that point. Lump sums withdrawn over your remaining allowances will be subject to Income Tax.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "Your LSA is £268,275 and your LSDBA is £1.0731 million.")
        
        doc.add_paragraph()
        mpaa_statement = self.data.get("user_input.mpaa_statement")
        p = doc.add_paragraph()
        run = p.add_run(mpaa_statement)
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        run.font.bold = True
    
    def _generate_part2(self, doc):
        """Generate Part 2 - My Recommendation."""
        print("  Generating Part 2...")
        
        part2 = self.template.get("part2", {})
        
        doc.add_page_break()
        self._add_heading1(doc, part2.get("heading1", "Part 2 - My Recommendation"))
        
        doc.add_paragraph()
        self._add_heading2(doc, part2.get("heading2", "Recommended Actions"))
        
        for item in part2.get("content", []):
            self._process_content_item(doc, item)
        
        # Employer Schemes section
        doc.add_paragraph()
        self._add_heading2(doc, "Employer Schemes")
        
        employer = self.data.sources.get("user_input", {}).get("employer_scheme", {})
        self._add_paragraph(doc, f"Your employer offers you the benefit of contribution matching for your current scheme. {employer.get('contribution_matching_status', '')}")
        
        doc.add_paragraph()
        
        # Employer scheme table
        self._add_employer_scheme_table(doc)
        
        doc.add_paragraph()
        
        # Employer comparison text
        cyc = self.data.sources.get("cyc", {}).get("employer_scheme_comparison", {})
        comparison_text = (
            f"Your employer's scheme has charges of {cyc.get('employer_scheme_charge', '')} per annum "
            f"compared to the charges on your St. James's Place Retirement Account of {cyc.get('sjp_charge', '')} per annum. "
            f"This means the recommended plan would need to grow by {cyc.get('outperformance_required', '')} more per annum, "
            f"or {cyc.get('monetary_equivalent_first_year', '')} more in the first year, than your employer's scheme "
            f"in order to match the fund you could have had if you had invested via your employer scheme."
        )
        self._add_paragraph(doc, comparison_text)
        
        doc.add_paragraph()
        
        taking_into = f"Taking into account my calculations, I have recommended you contribute to a St. James's Place Retirement Account rather than your employer's scheme because {employer.get('sjp_vs_employer_reason', '')}"
        self._add_paragraph(doc, taking_into)
        
        # Salary Sacrifice
        doc.add_paragraph()
        self._add_heading2(doc, "Salary Sacrifice")
        salary_sacrifice = self.data.get("illustration.salary_sacrifice.description")
        self._add_paragraph(doc, salary_sacrifice)
        
        # Legacy Preservation Trust
        doc.add_paragraph()
        self._add_heading2(doc, "Use of the Legacy Preservation Trust")
        self._add_paragraph(doc, "I have recommended you set up a Legacy Preservation Trust, a provision which we offer as part of our standard service and at no additional cost to you. It is specifically designed to receive death benefits, including accrued pension funds, from most pension schemes. If used on death, the trust will protect your pension benefits for your beneficiaries and will prevent the value of the pension forming part of their individual estates for Inheritance Tax purposes.")
        doc.add_paragraph()
        self._add_paragraph(doc, "While you accepted my recommendation, you have not yet completed the necessary paperwork as you are still deciding on how to structure the Trust. I left the relevant paperwork with you, and you will look to complete this as soon as possible as you understand until this has been completed, the Trust will not be created.")
        doc.add_paragraph()
        self._add_paragraph(doc, "I have also recommended you complete an Expression of Wish form, which you have agreed to. An Expression of Wish form ensures that your wishes are considered prior to payment of your pension death benefits. Please provide your completed form as soon as you can, and I will update your record.")
    
    def _generate_part3(self, doc):
        """Generate Part 3 - Impact of Replacement."""
        print("  Generating Part 3...")
        
        part3 = self.template.get("part3", {})
        
        doc.add_page_break()
        self._add_heading1(doc, part3.get("heading1", "Part 3 - Impact of Replacement"))
        
        doc.add_paragraph()
        self._add_heading2(doc, part3.get("heading2", "Pension Plans to be transferred"))
        
        # Transfer table
        self._add_transfer_impact_table(doc)
        
        doc.add_paragraph()
        
        # Plan name
        plan_name = self.data.get("cfr.pension_arrangements.scottish_widows.provider_plan")
        p = doc.add_paragraph()
        run = p.add_run(plan_name)
        run.font.bold = True
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        
        # Tax-Free Cash section
        doc.add_paragraph()
        self._add_heading2(doc, "Changes to Your Tax-Free Cash Entitlement")
        self._add_paragraph(doc, "There will be no change to your tax-free cash entitlement from the uncrystallised funds relating to the transfer of your plan from Scottish Widows.")
        
        # Impact of Charges section
        doc.add_paragraph()
        self._add_heading2(doc, "The Impact of Charges on your Benefits")
        self._add_paragraph(doc, "Moving funds as recommended into a Retirement Account will usually mean there will be a change in the level of charges which you pay. As part of my recommendation, I have therefore considered the difference in charges between your existing arrangements and the Retirement Account and set out the results in the table below.")
        
        doc.add_paragraph()
        
        spread_period = self.data.get("cyc.charges_spread_period")
        self._add_paragraph(doc, f"Our comparison includes the product charge, fund charge, ongoing advice charge and initial charge (as detailed in appendix). For our calculations, the initial charge has been spread over {spread_period}.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "The outperformance shown below is the additional growth your Retirement Account needs to achieve to match these funds, based on the increase in charges. If this level of outperformance is not achieved, your fund and therefore the benefits it can pay, will be lower than if you had not transferred.")
        
        doc.add_paragraph()
        
        # Outperformance table
        self._add_outperformance_table(doc)
        
        # Disadvantages section
        doc.add_paragraph()
        self._add_heading2(doc, "Potential Disadvantages of Replacing Your Existing Plan")
        self._add_paragraph(doc, "In addition to the above points the proposed Retirement Account has the following disadvantages:")
        
        ceding = self.data.sources.get("ceding_info", {}).get("disadvantages", {})
        cyc_disadvantages = self.data.sources.get("cyc", {}).get("disadvantages", {})
        
        self._add_bullet(doc, ceding.get("higher_charges", "Higher charges"))
        self._add_bullet(doc, cyc_disadvantages.get("new_initial_advice_charge", "A new Initial Advice Charge will be applied"))
        self._add_bullet(doc, ceding.get("fewer_funds", "Fewer funds to invest in"))
        self._add_bullet(doc, ceding.get("loss_lifestyling", "Loss of Lifestyling"))
        self._add_bullet(doc, ceding.get("loss_death_cover", "Loss of accidental death cover"))
        
        # Special terms section
        doc.add_paragraph()
        self._add_heading2(doc, "Special terms")
        
        special_terms = self.data.sources.get("cyc", {}).get("special_terms", {})
        self._add_paragraph(doc, "Following the calculation of the impact of the difference in charges between the recommended St. James's Place Retirement Account and your existing arrangements, I have arranged Special Terms to reduce the Initial Advice Charge payable on transfer to the St. James's Place Retirement Account. These are reflected in the reduced charges shown on the Illustration that accompanies this report.")
        
        doc.add_paragraph()
        
        combined_text = f"After applying these terms, the combined outperformance is {special_terms.get('combined_outperformance', '')} per annum which is equivalent to {special_terms.get('monetary_equivalent', '')} in the next twelve months."
        self._add_paragraph(doc, combined_text)
        
        doc.add_paragraph()
        self._add_paragraph(doc, "You understand past performance is no guide to future performance and there is no guarantee St. James's Place funds will outperform an existing provider. Having discussed this in detail, you felt there was a reasonable opportunity for sufficient growth to be achieved and you are willing to accept the risk it might not be.")
    
    def _generate_part4(self, doc):
        """Generate Part 4 - Attitude to Risk."""
        print("  Generating Part 4...")
        
        part4 = self.template.get("part4", {})
        
        doc.add_page_break()
        self._add_heading1(doc, part4.get("heading1", "Part 4 - Your Attitude to Risk and fund selection"))
        
        doc.add_paragraph()
        self._add_paragraph(doc, "We had a conversation about investment risk as part of our discussions. Some key factors we discussed were your objectives, your investment experience, the time horizon over which you are investing and your attitude to, and ability to withstand, investment losses.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "We also discussed the range of example portfolios and funds offered by St. James's Place, and the importance of holding a diversified range of investments.")
        
        # Investment knowledge
        doc.add_paragraph()
        self._add_heading2(doc, "Your investment knowledge and experience")
        self._add_paragraph(doc, "When considering your level of investment knowledge and experience, a key factor is the extent to which your existing or previous investments have been exposed to market volatility.")
        
        doc.add_paragraph()
        # Intro text
        investment_intro = self.data.get("cfr.client_profile.investment_knowledge_intro")
        self._add_paragraph(doc, investment_intro)
        
        # Bullet points
        investment_bullets = self.data.get("cfr.client_profile.investment_knowledge_bullets")
        if isinstance(investment_bullets, list):
            for bullet in investment_bullets:
                self._add_bullet(doc, bullet)
        
        # Capacity for loss
        doc.add_paragraph()
        self._add_heading2(doc, "Your capacity for loss")
        self._add_paragraph(doc, "Capacity for loss relates to your ability to withstand investment losses and cope financially should there be a market downturn either at the point you need to access your investment, or if investment returns are less favourable than expected. Considering your capacity for loss, in conjunction with your investment knowledge and experience, is particularly important when making investment decisions.")
        
        doc.add_paragraph()
        # Intro text
        capacity_intro = self.data.get("cfr.client_profile.capacity_for_loss_intro")
        self._add_paragraph(doc, capacity_intro)
        
        # Bullet points
        capacity_bullets = self.data.get("cfr.client_profile.capacity_for_loss_bullets")
        if isinstance(capacity_bullets, list):
            for bullet in capacity_bullets:
                self._add_bullet(doc, bullet)
        
        doc.add_paragraph()
        emergency_fund = self.data.get("cfr.client_profile.emergency_fund")
        self._add_paragraph(doc, f"An emergency fund of {emergency_fund} is held in accessible cash accounts.")
        
        doc.add_paragraph()
        emergency_reason = self.data.get("user_input.client_profile.emergency_fund_reason")
        self._add_paragraph(doc, f"I believe that your emergency fund is sufficient for your needs because {emergency_reason}")
        
        # Risk Profile
        doc.add_paragraph()
        self._add_heading2(doc, "Your Risk Profile")
        
        risk_profile = self.data.get("user_input.client_profile.risk_profile")
        risk_text = f"Taking into account all these factors, we agreed you are a {risk_profile} Risk investor on our risk spectrum. Investors in this category are willing to take a balanced approach to investment risk. Such investors would like the value of their money to keep pace with inflation and are prepared to invest over a minimum of five years. However, neither do they want to take too much investment risk. Therefore, any investment in this category will be invested, either with a more balanced weighting between equities, bonds and gilts, or a higher weighting to equities. This is to enable moderate growth over the longer term, compared to investment solely into cash, bonds or gilts. Such investors can accept moderate to significant falls in the value of their investment and also understand it is possible to lose some of their initial money (capital)."
        self._add_paragraph(doc, risk_text)
        
        doc.add_paragraph()
        self._add_paragraph(doc, "I recommended that you invest in a fund selection as follows:")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "Contribution: Transfer")
        
        # Fund selection table
        self._add_fund_selection_table(doc)
        
        doc.add_paragraph()
        
        # Fund recommendation reason
        fund_reason = self.data.get("user_input.fund_selection.recommendation_reason")
        self._add_paragraph(doc, fund_reason)
        
        # Equity content note
        equity_content = self.data.get("illustration.fund_selection.overall_equity_content")
        self._add_paragraph(doc, equity_content)
        
        doc.add_paragraph()
        
        # Additional fund notes
        additional_notes = self.data.get("user_input.fund_selection.additional_fund_notes")
        self._add_paragraph(doc, additional_notes)
    
    def _generate_appendix_i(self, doc):
        """Generate Appendix I - Further details to the background of the recommendation made."""
        print("  Generating Appendix I...")
        
        doc.add_page_break()
        self._add_heading1(doc, "Appendix i - Further details to the background of the recommendation made")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "This Appendix summarises further factual information which I learnt during our discussions and acts as a useful aid to explaining the reasons for my recommendations.")
        
        # Liabilities
        doc.add_paragraph()
        self._add_heading2(doc, "Liabilities")
        debts = self.data.get("user_input.client_profile.debts")
        self._add_paragraph(doc, f"There were {debts} to consider as part of my recommendation")
        
        # Alternatives - Existing Provider
        doc.add_paragraph()
        self._add_heading2(doc, "Alternatives Available – Existing Provider(s)")
        self._add_paragraph(doc, "The plan being transferred could have been left as they are now, and an attempt made to resolve any concerns directly with the existing providers. This would avoid any increase in charges on transfer.")
        
        doc.add_paragraph()
        alternative_reason = self.data.get("user_input.client_profile.alternative_not_recommended_reason")
        self._add_paragraph(doc, f"Having considered this option and the reasons for wanting to transfer, this was not recommended because {alternative_reason}")
        
        # Alternatives - Stakeholder Scheme
        doc.add_paragraph()
        self._add_heading2(doc, "Alternatives Available – Stakeholder Scheme")
        self._add_paragraph(doc, "The rules for a Stakeholder plan do not allow advice charges to be paid via the plan meaning these charges have to be paid separately from your taxed income. As a result of this, the overall cost difference is typically not material, depending on your fund selection. A recommendation for a St. James's Place Retirement Account allows access to the St. James's Place Approach to Investment Management at essentially no additional cost. A contribution into a Stakeholder plan has therefore not been considered.")
        
        # Death Benefits
        doc.add_paragraph()
        self._add_heading2(doc, "Death Benefits")
        self._add_paragraph(doc, "On your death, your beneficiaries can take benefits from your pension fund as either; an income, a lump sum paid to them or paid to a trust, or a combination of these.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "If you die before age 75, your pension death benefits can usually be paid tax free, as long as benefits are taken within two years of the scheme being notified of your death, and any relevant lump sum payments are within your remaining Lump Sum and Death Benefit Allowance (LSDBA). Relevant lump sum payments over your LSDBA that are paid to a beneficiary will be subject to Income Tax payable at their highest marginal rate, and if paid to a discretionary trust, tax will be payable at 45%.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "If you die after you reach age 75, any benefits paid to your beneficiaries will be subject to Income Tax payable at their highest marginal rate. If benefits are paid to a discretionary trust, they will be subject to a 45% tax charge.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "The October 2024 Budget announced that from April 2027, on death, most pension funds will form part of the estate for Inheritance Tax purposes, regardless of whether the benefit is paid to an individual or trust. Legislation is currently being drafted, so it is subject to change. As part of our ongoing reviews, we will consider the impact on your estate and any appropriate action required.")
        
        # Further Relevant Client Specific Information
        doc.add_paragraph()
        self._add_heading2(doc, "Further Relevant Client Specific Information")
        
        annual_income = self.data.get("cfr.client_profile.annual_income")
        tax_bracket = self.data.get("cfr.client_profile.tax_bracket")
        self._add_paragraph(doc, f"You have an income of {annual_income} per annum which falls into the {tax_bracket} tax bracket.")
        
        doc.add_paragraph()
        has_will = self.data.get("cfr.client_profile.has_will")
        self._add_paragraph(doc, f"{has_will}, so I recommended you seek legal advice to review your Will arrangements.")
        
        doc.add_paragraph()
        self._add_paragraph(doc, "I recommend you seek legal advice to arrange a Power of Attorney to ensure someone you trust can manage your personal or financial affairs if you lose capacity.")
        
        # Determining ATR
        doc.add_paragraph()
        self._add_heading2(doc, "Determining Your Attitude to Risk (ATR)")
        self._add_paragraph(doc, "We had a conversation about investment risk as part of our discussions. Some key factors we discussed were your objectives, your investment experience, the time horizon over which you are investing and your attitude to, and ability to withstand, investment losses.")
        
        doc.add_paragraph()
        investment_aim = self.data.get("user_input.client_profile.investment_aim")
        self._add_paragraph(doc, f"You intend to use your investment to {investment_aim}.")
        
        doc.add_paragraph()
        timeframe = self.data.get("illustration.investment_timeframe")
        self._add_paragraph(doc, f"With this aim in mind, the timeframe for your investment into your St. James's Place Retirement Account is {timeframe}")
    
    def _generate_appendix(self, doc):
        """Generate the Appendix section."""
        print("  Generating Appendix...")
        
        # Create landscape section
        section = doc.add_section()
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.orientation = 1  # Landscape
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        appendix = self.template.get("appendix", {})
        self._add_heading1(doc, appendix.get("title", "Appendix ii - Product Comparison"))
        
        doc.add_paragraph()
        
        # Product comparison table
        self._add_product_comparison_table(doc)
    
    def _process_content_item(self, doc, item):
        """Process a content item from the template."""
        item_type = item.get("type", "paragraph")
        
        if item_type == "paragraph":
            text = item.get("text", "")
            resolved = self.data.resolve_placeholders(text)
            doc.add_paragraph()
            self._add_paragraph(doc, resolved)
        
        elif item_type == "heading2":
            text = item.get("text", "")
            resolved = self.data.resolve_placeholders(text)
            doc.add_paragraph()
            self._add_heading2(doc, resolved)
        
        elif item_type == "bullet_list":
            source = item.get("source", "")
            doc.add_paragraph()
            
            if source:
                # Get bullets from data source
                data = self.data.get(source)
                if isinstance(data, list):
                    for bullet_text in data:
                        self._add_bullet(doc, bullet_text)
                elif isinstance(data, dict):
                    keys = item.get("keys", list(data.keys()))
                    for key in keys:
                        if key in data:
                            self._add_bullet(doc, data[key])
        
        elif item_type == "table":
            table_id = item.get("table_id", "")
            doc.add_paragraph()
            if table_id == "pension_arrangements":
                self._add_pension_arrangements_table(doc)
            elif table_id == "recommendation_table":
                self._add_recommendation_table(doc)
    
    # ========================================================================
    # HELPER METHODS
    # ========================================================================
    
    def _add_paragraph(self, doc, text):
        """Add a paragraph with default styling."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        return p
    
    def _add_heading1(self, doc, text):
        """Add a Heading 1 paragraph."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Noe Display SJP Bold"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3C)
        p.paragraph_format.space_after = Pt(6)
        return p
    
    def _add_heading2(self, doc, text):
        """Add a Heading 2 paragraph."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Poppins SemiBold"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3C)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        return p
    
    def _add_bullet(self, doc, text):
        """Add a bullet point paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(f"•\t{text}")
        run.font.name = "Poppins"
        run.font.size = Pt(10)
        return p
    
    def _apply_style(self, paragraph, style_config):
        """Apply style configuration to a paragraph."""
        for run in paragraph.runs:
            if style_config.get("bold"):
                run.font.bold = True
            run.font.name = style_config.get("font_family", "Poppins")
            run.font.size = Pt(style_config.get("font_size", 10))
    
    def _add_page_numbers(self, doc):
        """Add page numbers to footer."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run()
            
            # PAGE field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run.font.name = 'Poppins'
            run.font.size = Pt(10)
    
    # ========================================================================
    # TABLE METHODS
    # ========================================================================
    
    def _add_pension_arrangements_table(self, doc):
        """Add the pension arrangements table."""
        columns = ["Provider Type of Plan / Plan Number", "Current Value", "Transfer Value", "Ongoing Contribution", "Proposed Action"]
        
        cfr = self.data.sources.get("cfr", {}).get("pension_arrangements", {})
        user_input = self.data.sources.get("user_input", {}).get("part1_objectives", {})
        illustration = self.data.sources.get("illustration", {}).get("aviva_plan", {})
        
        rows = []
        
        # SJP row
        sjp = cfr.get("sjp_retirement_account", {})
        rows.append([
            sjp.get("provider_plan", ""),
            sjp.get("current_value", ""),
            sjp.get("transfer_value", ""),
            sjp.get("ongoing_contribution", ""),
            sjp.get("proposed_action", "")
        ])
        
        # Scottish Widows row
        sw = cfr.get("scottish_widows", {})
        rows.append([
            sw.get("provider_plan", ""),
            sw.get("current_value", ""),
            sw.get("transfer_value", ""),
            sw.get("ongoing_contribution", ""),
            user_input.get("scottish_widows_proposed_action", "Transfer")
        ])
        
        # Aviva row
        aviva = cfr.get("aviva", {})
        rows.append([
            aviva.get("provider_plan", ""),
            illustration.get("current_value", ""),
            illustration.get("transfer_value", ""),
            illustration.get("ongoing_contribution", ""),
            user_input.get("aviva_proposed_action", "None")
        ])
        
        self._create_table(doc, columns, rows)
    
    def _add_recommendation_table(self, doc):
        """Add the recommendation table."""
        columns = ["Source", "Transfer/Redirection", "Lump Sum Amount", "Regular Contribution", "Indexation"]
        
        cfr = self.data.sources.get("cfr", {}).get("part2_recommendation", {})
        user_input = self.data.sources.get("user_input", {}).get("part2_recommendation", {})
        
        rows = [[
            cfr.get("source_plan", ""),
            "Transfer",
            cfr.get("lump_sum_amount", ""),
            user_input.get("regular_contribution", "None"),
            "None"
        ]]
        
        self._create_table(doc, columns, rows)
    
    def _add_employer_scheme_table(self, doc):
        """Add the employer scheme table."""
        columns = ["Contribution Type", "Source", "Can the employer accept this type of contribution?", "Type of Employer Scheme"]
        
        employer = self.data.sources.get("user_input", {}).get("employer_scheme", {})
        
        rows = [[
            employer.get("contribution_type", ""),
            employer.get("source", ""),
            employer.get("can_accept", ""),
            employer.get("scheme_type", "")
        ]]
        
        self._create_table(doc, columns, rows)
    
    def _add_transfer_impact_table(self, doc):
        """Add the transfer impact table."""
        columns = ["Provider Plan", "Transfer Value", "Charge on Transfer", "Protected Tax-Free Cash", "Guarantees", "Terminal Bonus"]
        
        cfr = self.data.sources.get("cfr", {}).get("part3_impact", {})
        ceding = self.data.sources.get("ceding_info", {}).get("part3_transfer_details", {})
        
        rows = [[
            cfr.get("provider_plan", ""),
            cfr.get("transfer_value", ""),
            ceding.get("charge_on_transfer", ""),
            ceding.get("protected_tfc", ""),
            ceding.get("guarantees", ""),
            ceding.get("terminal_bonus", "")
        ]]
        
        self._create_table(doc, columns, rows)
    
    def _add_outperformance_table(self, doc):
        """Add the outperformance table."""
        columns = ["Provider Plan", "Action", "Level of Outperformance Required", "Monetary Equivalent in First Year"]
        
        cfr = self.data.sources.get("cfr", {}).get("part3_impact", {})
        cyc = self.data.sources.get("cyc", {}).get("outperformance_table", {}).get("scottish_widows", {})
        
        rows = [[
            cfr.get("provider_plan", ""),
            "Transfer",
            cyc.get("level_of_outperformance", ""),
            cyc.get("monetary_equivalent", "")
        ]]
        
        self._create_table(doc, columns, rows)
    
    def _add_fund_selection_table(self, doc):
        """Add the fund selection table."""
        columns = ["Category", "Fund Selection", "Fund Type", "Risk Profile", "Percentage of investment"]
        
        funds = self.data.sources.get("user_input", {}).get("fund_selection", {}).get("table", [])
        
        rows = []
        for fund in funds:
            rows.append([
                fund.get("category", ""),
                fund.get("fund_selection", ""),
                fund.get("fund_type", ""),
                fund.get("risk_profile", ""),
                fund.get("percentage", "")
            ])
        
        self._create_table(doc, columns, rows)
    
    def _add_product_comparison_table(self, doc):
        """Add the product comparison table in appendix."""
        columns = [
            "Provider Plan Name / Number",
            "Initial Advice Charge",
            "Ongoing Advice Charge",
            "Ongoing Product Charge",
            "Ongoing Fund Charge",
            "Annual Management Charge",
            "Early Withdrawal Charges",
            "Allocation Rate",
            "Total External Management Charge",
            "Plan Charge",
            "Number of Funds",
            "Maximum Number of Funds",
            "Any Guarantees Applicable"
        ]
        
        cyc = self.data.sources.get("cyc", {}).get("product_comparison", {}).get("sjp_retirement_account", {})
        user_input = self.data.sources.get("user_input", {}).get("product_comparison", {})
        ceding = self.data.sources.get("ceding_info", {}).get("product_comparison_scottish_widows", {})
        
        rows = [
            # SJP row
            [
                "St. James's Place Retirement Account",
                cyc.get("initial_advice_charge", ""),
                cyc.get("ongoing_advice_charge", ""),
                cyc.get("ongoing_product_charge", ""),
                cyc.get("ongoing_fund_charge", ""),
                "None",
                "None",
                user_input.get("allocation_rate", ""),
                "-",
                "None",
                "43",
                "Unlimited",
                "None"
            ],
            # Scottish Widows row
            [
                self.data.get("cfr.pension_arrangements.scottish_widows.provider_plan"),
                ceding.get("initial_advice_charge", ""),
                ceding.get("ongoing_advice_charge", ""),
                ceding.get("ongoing_product_charge", ""),
                ceding.get("ongoing_fund_charge", ""),
                ceding.get("annual_management_charge", ""),
                ceding.get("early_withdrawal_charges", ""),
                ceding.get("allocation_rate", ""),
                ceding.get("total_external_charge", ""),
                ceding.get("plan_charge", ""),
                ceding.get("number_of_funds", ""),
                ceding.get("max_funds", ""),
                ceding.get("guarantees", "")
            ]
        ]
        
        self._create_table(doc, columns, rows, header_vertical=True)
    
    def _create_table(self, doc, columns, rows, header_vertical=False):
        """Create a styled table."""
        table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
        table.style = 'Table Grid'
        table.autofit = False
        
        # Header row
        header_row = table.rows[0]
        for i, col_name in enumerate(columns):
            cell = header_row.cells[i]
            cell.text = col_name
            
            # Turquoise background
            set_cell_background(cell, "3FDCC8")
            
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Poppins"
                    run.font.size = Pt(9)
                    run.font.bold = True
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            if header_vertical:
                set_vertical_text(cell)
        
        # Data rows
        for row_idx, row_data in enumerate(rows, 1):
            data_row = table.rows[row_idx]
            for col_idx, cell_value in enumerate(row_data):
                cell = data_row.cells[col_idx]
                cell.text = str(cell_value) if cell_value else ""
                
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = "Poppins"
                        run.font.size = Pt(9)
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def main():
    """Main entry point."""
    print("\n" + "="*60)
    print("TEMPLATE DOCUMENT GENERATOR")
    print("="*60)
    
    # Generate document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"generated_documents/template_output_{timestamp}.docx"
    
    generator = TemplateDocumentGenerator()
    generator.generate(output_file)


if __name__ == "__main__":
    main()
