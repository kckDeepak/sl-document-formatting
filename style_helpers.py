"""
Style Helper Module for Word Document Generation
Provides reusable functions for applying JSON-driven styles to Word elements.

Values updated based on XML extraction from client's ideal document.
"""

from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap


# ============================================================================
# EXACT VALUES FROM XML EXTRACTION (ideal_xml_analysis.json)
# ============================================================================
# line="240" with lineRule="auto" = 1.0 (single spacing)
# left="714" twips = 35.7 pt = 0.4958 inches
# hanging="357" twips = 17.85 pt = 0.2479 inches
# All paragraphs have after="0" (no space after)
# Alignment is "left" (not "both"/justify)

# ============================================================================
# DEFAULT STYLES - Updated with exact XML values
# ============================================================================

DEFAULT_STYLES = {
    "document": {
        "font_family": "Poppins",
        "font_size": 10,
        "font_color": "000000",
        "line_spacing": 1.0,  # XML: line="240", lineRule="auto"
        "line_spacing_rule": "SINGLE",
        "space_before": 0,
        "space_after": 0,  # XML: after="0"
        "alignment": "LEFT"  # XML: jc="left"
    },
    # MAIN HEADINGS - Part 1, Part 2, Part 3, Part 4, Part 5, Appendix i, ii
    "heading1": {
        "font_family": "Noe Display SJP Bold",  # From ideal document
        "font_size": 20,  # 20pt as per ideal
        "font_color": "0F1E3C",  # Dark navy from ideal
        "bold": False,  # Font itself is bold
        "space_before": 0,
        "space_after": 6,  # 6pt after as per ideal
        "alignment": "LEFT",
        "page_break_before": True,
        "keep_with_next": True
    },
    # SUB HEADINGS - Section titles within parts
    "heading2": {
        "font_family": "Poppins SemiBold",  # From ideal document
        "font_size": 12,  # 12pt as per ideal
        "font_color": "0F1E3C",  # Dark navy from ideal
        "bold": False,  # Font itself is semibold
        "space_before": 6,
        "space_after": 6,
        "line_spacing": 1.0,
        "alignment": "LEFT",
        "keep_with_next": True
    },
    # MINOR HEADINGS - Smaller section titles
    "heading3": {
        "font_family": "Poppins SemiBold",  # From ideal document
        "font_size": 11,  # 11pt as per ideal
        "font_color": "0F1E3C",  # Dark navy from ideal
        "bold": False,
        "space_before": 0,
        "space_after": 0,
        "alignment": "LEFT",
        "keep_with_next": True
    },
    "bullet": {
        "font_family": "Poppins",
        "font_size": 10,
        "font_color": "000000",
        "line_spacing": 1.0,
        "space_before": 0,
        "space_after": 0,
        "alignment": "LEFT",
        "left_indent": 0.4958,  # inches
        "hanging_indent": 0.2479  # inches
    },
    # TABLE HEADER - Turquoise background with black text
    "table_header": {
        "font_family": "Poppins",
        "font_size": 8,
        "font_color": "000000",  # Black text on turquoise
        "bold": True,
        "alignment": "CENTER",
        "vertical_alignment": "CENTER",
        "background_color": "3FDCC8"  # Turquoise from ideal document
    },
    "table_cell": {
        "font_family": "Poppins",
        "font_size": 8,
        "font_color": "000000",
        "alignment": "CENTER",
        "vertical_alignment": "CENTER"
    },
    "footnote": {
        "font_family": "Poppins",
        "font_size": 9,
        "font_color": "000000",
        "italic": True,
        "space_after": 0,
        "alignment": "LEFT"
    },
    "closing": {
        "font_family": "Poppins",
        "font_size": 10,
        "font_color": "000000",
        "space_before": 0,
        "space_after": 0,
        "alignment": "LEFT"
    }
}


# ============================================================================
# ALIGNMENT MAPPING
# ============================================================================

ALIGNMENT_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
}

VERTICAL_ALIGNMENT_MAP = {
    "TOP": WD_ALIGN_VERTICAL.TOP,
    "CENTER": WD_ALIGN_VERTICAL.CENTER,
    "BOTTOM": WD_ALIGN_VERTICAL.BOTTOM
}

LINE_SPACING_RULE_MAP = {
    "SINGLE": WD_LINE_SPACING.SINGLE,
    "ONE_POINT_FIVE": WD_LINE_SPACING.ONE_POINT_FIVE,
    "DOUBLE": WD_LINE_SPACING.DOUBLE,
    "MULTIPLE": WD_LINE_SPACING.MULTIPLE,
    "EXACTLY": WD_LINE_SPACING.EXACTLY,
    "AT_LEAST": WD_LINE_SPACING.AT_LEAST
}


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def hex_to_rgb(hex_color):
    """Convert hex color string to RGBColor"""
    hex_color = hex_color.lstrip('#')
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


def get_style_config(style_name, overrides=None):
    """
    Get style configuration with optional overrides.
    
    Args:
        style_name: Name of the default style (e.g., 'document', 'heading1')
        overrides: Dict of style properties to override defaults
    
    Returns:
        Merged style configuration dict
    """
    base_style = DEFAULT_STYLES.get(style_name, DEFAULT_STYLES['document']).copy()
    if overrides:
        base_style.update(overrides)
    return base_style


def apply_run_style(run, style_config):
    """
    Apply styling to a run (text fragment).
    
    Args:
        run: python-docx Run object
        style_config: Dict with styling properties
    """
    # Font family
    if 'font_family' in style_config:
        run.font.name = style_config['font_family']
        # Set font for complex scripts as well
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), style_config['font_family'])
        rFonts.set(qn('w:hAnsi'), style_config['font_family'])
        rFonts.set(qn('w:cs'), style_config['font_family'])
    
    # Font size
    if 'font_size' in style_config:
        run.font.size = Pt(style_config['font_size'])
    
    # Font color
    if 'font_color' in style_config:
        run.font.color.rgb = hex_to_rgb(style_config['font_color'])
    
    # Bold
    if 'bold' in style_config:
        run.font.bold = style_config['bold']
    
    # Italic
    if 'italic' in style_config:
        run.font.italic = style_config['italic']
    
    # Underline
    if 'underline' in style_config:
        run.font.underline = style_config['underline']
    
    # Superscript
    if 'superscript' in style_config:
        run.font.superscript = style_config['superscript']
    
    # Subscript
    if 'subscript' in style_config:
        run.font.subscript = style_config['subscript']
    
    # Character spacing
    if 'char_spacing' in style_config:
        run.font.spacing = Pt(style_config['char_spacing'])


def apply_paragraph_style(paragraph, style_config):
    """
    Apply styling to a paragraph.
    
    Args:
        paragraph: python-docx Paragraph object
        style_config: Dict with styling properties
    """
    pf = paragraph.paragraph_format
    
    # Alignment
    if 'alignment' in style_config:
        alignment = style_config['alignment']
        if isinstance(alignment, str):
            paragraph.alignment = ALIGNMENT_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
        else:
            paragraph.alignment = alignment
    
    # Line spacing
    if 'line_spacing' in style_config:
        line_spacing = style_config['line_spacing']
        rule = style_config.get('line_spacing_rule', 'MULTIPLE')
        
        if rule == 'MULTIPLE':
            pf.line_spacing = line_spacing
        elif rule == 'EXACTLY':
            pf.line_spacing = Pt(line_spacing)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif rule == 'AT_LEAST':
            pf.line_spacing = Pt(line_spacing)
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    
    # Space before
    if 'space_before' in style_config:
        pf.space_before = Pt(style_config['space_before'])
    
    # Space after
    if 'space_after' in style_config:
        pf.space_after = Pt(style_config['space_after'])
    
    # Left indent
    if 'left_indent' in style_config:
        pf.left_indent = Inches(style_config['left_indent'])
    
    # Right indent
    if 'right_indent' in style_config:
        pf.right_indent = Inches(style_config['right_indent'])
    
    # First line indent (can be negative for hanging)
    if 'first_line_indent' in style_config:
        pf.first_line_indent = Inches(style_config['first_line_indent'])
    
    # Hanging indent
    if 'hanging_indent' in style_config:
        pf.first_line_indent = Inches(-style_config['hanging_indent'])
    
    # Keep with next
    if 'keep_with_next' in style_config:
        pf.keep_with_next = style_config['keep_with_next']
    
    # Keep together
    if 'keep_together' in style_config:
        pf.keep_together = style_config['keep_together']
    
    # Page break before
    if 'page_break_before' in style_config:
        pf.page_break_before = style_config['page_break_before']
    
    # Apply run styling to all runs in paragraph
    for run in paragraph.runs:
        apply_run_style(run, style_config)


def apply_full_paragraph_style(paragraph, text, style_config):
    """
    Clear paragraph, add text, and apply full styling.
    
    Args:
        paragraph: python-docx Paragraph object
        text: Text content to add
        style_config: Dict with styling properties
    
    Returns:
        The created run
    """
    paragraph.clear()
    run = paragraph.add_run(text)
    apply_run_style(run, style_config)
    apply_paragraph_style(paragraph, style_config)
    return run


def create_styled_paragraph(document, text, style_name, style_overrides=None):
    """
    Create a new paragraph with full styling.
    
    Args:
        document: python-docx Document object
        text: Text content
        style_name: Name of default style to use
        style_overrides: Optional dict to override default style
    
    Returns:
        The created paragraph
    """
    style_config = get_style_config(style_name, style_overrides)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_run_style(run, style_config)
    apply_paragraph_style(paragraph, style_config)
    return paragraph


# ============================================================================
# CELL STYLING FUNCTIONS
# ============================================================================

def set_cell_background(cell, color_hex):
    """Set background/shading color for a table cell"""
    color_hex = color_hex.lstrip('#')
    cell_pr = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color_hex)
    cell_shading.set(qn('w:val'), 'clear')
    cell_pr.append(cell_shading)


def set_vertical_text(cell):
    """Set text direction to vertical (bottom to top) in a cell"""
    tc_pr = cell._element.get_or_add_tcPr()
    text_direction = OxmlElement('w:textDirection')
    text_direction.set(qn('w:val'), 'btLr')
    tc_pr.append(text_direction)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins/padding"""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    
    for margin_name, margin_val in [('top', top), ('bottom', bottom), 
                                      ('left', left), ('right', right)]:
        margin_elem = OxmlElement(f'w:{margin_name}')
        margin_elem.set(qn('w:w'), str(int(margin_val * 20)))  # Convert to twips
        margin_elem.set(qn('w:type'), 'dxa')
        tc_mar.append(margin_elem)
    
    tc_pr.append(tc_mar)


def apply_cell_style(cell, style_config):
    """
    Apply full styling to a table cell.
    
    Args:
        cell: python-docx table cell
        style_config: Dict with styling properties
    """
    # Background color
    if 'background_color' in style_config:
        set_cell_background(cell, style_config['background_color'])
    
    # Vertical text
    if style_config.get('vertical_text', False):
        set_vertical_text(cell)
    
    # Vertical alignment
    if 'vertical_alignment' in style_config:
        va = style_config['vertical_alignment']
        if isinstance(va, str):
            cell.vertical_alignment = VERTICAL_ALIGNMENT_MAP.get(va.upper(), WD_ALIGN_VERTICAL.CENTER)
        else:
            cell.vertical_alignment = va
    
    # Apply paragraph styling to cell content
    if cell.paragraphs:
        for para in cell.paragraphs:
            apply_paragraph_style(para, style_config)


def style_cell_text(cell, text, style_config):
    """
    Set cell text with full styling.
    
    Args:
        cell: python-docx table cell
        text: Text content
        style_config: Dict with styling properties
    """
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(text)
    apply_run_style(run, style_config)
    apply_paragraph_style(para, style_config)
    apply_cell_style(cell, style_config)


# ============================================================================
# SUPERSCRIPT HANDLING
# ============================================================================

SUPERSCRIPT_MAP = {'¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5', 
                   '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9', '⁰': '0'}

def has_superscript_chars(text):
    """Check if text contains superscript characters"""
    return any(char in text for char in SUPERSCRIPT_MAP.keys())


def add_text_with_superscript(paragraph, text, style_config):
    """
    Add text to paragraph with proper superscript formatting.
    
    Args:
        paragraph: python-docx Paragraph object
        text: Text that may contain superscript characters (¹²³)
        style_config: Base style config for the text
    """
    current_pos = 0
    
    for i, char in enumerate(text):
        if char in SUPERSCRIPT_MAP:
            # Add text before superscript
            if i > current_pos:
                run = paragraph.add_run(text[current_pos:i])
                apply_run_style(run, style_config)
            
            # Add superscript character
            sup_config = style_config.copy()
            sup_config['superscript'] = True
            run = paragraph.add_run(SUPERSCRIPT_MAP[char])
            apply_run_style(run, sup_config)
            
            current_pos = i + 1
    
    # Add remaining text
    if current_pos < len(text):
        run = paragraph.add_run(text[current_pos:])
        apply_run_style(run, style_config)


# ============================================================================
# BULLET LIST FUNCTIONS
# ============================================================================

def set_bullet_style(paragraph, bullet_char='•', style_config=None):
    """
    Apply bullet styling to a paragraph.
    
    Args:
        paragraph: python-docx Paragraph object
        bullet_char: Character to use as bullet
        style_config: Style configuration dict
    """
    if style_config is None:
        style_config = get_style_config('bullet')
    
    # Set indentation for bullet
    pf = paragraph.paragraph_format
    
    left_indent = style_config.get('left_indent', 0.5)
    hanging = style_config.get('hanging_indent', 0.25)
    
    pf.left_indent = Inches(left_indent)
    pf.first_line_indent = Inches(-hanging)
    
    apply_paragraph_style(paragraph, style_config)


def create_bullet_paragraph(document, text, style_config=None, bullet_char='•'):
    """
    Create a bullet point paragraph.
    
    Args:
        document: python-docx Document object
        text: Bullet text content
        style_config: Optional style overrides
        bullet_char: Bullet character to use
    
    Returns:
        Created paragraph
    """
    if style_config is None:
        style_config = get_style_config('bullet')
    else:
        style_config = get_style_config('bullet', style_config)
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{bullet_char}\t{text}")
    apply_run_style(run, style_config)
    set_bullet_style(paragraph, bullet_char, style_config)
    
    return paragraph


# ============================================================================
# SECTION & PAGE FUNCTIONS
# ============================================================================

def apply_section_settings(section, settings):
    """
    Apply settings to a document section.
    
    Args:
        section: python-docx Section object
        settings: Dict with section properties
    """
    # Page size
    if 'page_width' in settings:
        width = settings['page_width']
        if isinstance(width, (int, float)):
            section.page_width = Inches(width)
        else:
            section.page_width = width
    
    if 'page_height' in settings:
        height = settings['page_height']
        if isinstance(height, (int, float)):
            section.page_height = Inches(height)
        else:
            section.page_height = height
    
    # Orientation (0=Portrait, 1=Landscape)
    if 'orientation' in settings:
        orientation = settings['orientation']
        if isinstance(orientation, str):
            section.orientation = 1 if orientation.upper() == 'LANDSCAPE' else 0
        else:
            section.orientation = orientation
    
    # Margins
    if 'top_margin' in settings:
        margin = settings['top_margin']
        section.top_margin = Cm(margin) if isinstance(margin, (int, float)) else margin
    
    if 'bottom_margin' in settings:
        margin = settings['bottom_margin']
        section.bottom_margin = Cm(margin) if isinstance(margin, (int, float)) else margin
    
    if 'left_margin' in settings:
        margin = settings['left_margin']
        section.left_margin = Cm(margin) if isinstance(margin, (int, float)) else margin
    
    if 'right_margin' in settings:
        margin = settings['right_margin']
        section.right_margin = Cm(margin) if isinstance(margin, (int, float)) else margin
    
    # Header/footer distance
    if 'header_distance' in settings:
        section.header_distance = Inches(settings['header_distance'])
    
    if 'footer_distance' in settings:
        section.footer_distance = Inches(settings['footer_distance'])


def create_landscape_section(document, settings=None):
    """
    Add a new landscape A4 section to document.
    
    Args:
        document: python-docx Document object
        settings: Optional dict to override default landscape settings
    
    Returns:
        The new section
    """
    default_settings = {
        'page_width': 11.69,  # A4 height in inches
        'page_height': 8.27,  # A4 width in inches
        'orientation': 'LANDSCAPE',
        'top_margin': 2.54,
        'bottom_margin': 1.27,
        'left_margin': 5.0,
        'right_margin': 5.0
    }
    
    if settings:
        default_settings.update(settings)
    
    section = document.add_section()
    apply_section_settings(section, default_settings)
    
    return section


def create_portrait_section(document, settings=None):
    """
    Add a new portrait A4 section to document.
    
    Args:
        document: python-docx Document object
        settings: Optional dict to override default portrait settings
    
    Returns:
        The new section
    """
    default_settings = {
        'page_width': 8.27,   # A4 width in inches
        'page_height': 11.69,  # A4 height in inches
        'orientation': 'PORTRAIT',
        'top_margin': 2.54,
        'bottom_margin': 2.54,
        'left_margin': 2.54,
        'right_margin': 2.54
    }
    
    if settings:
        default_settings.update(settings)
    
    section = document.add_section()
    apply_section_settings(section, default_settings)
    
    return section


# ============================================================================
# TABLE FUNCTIONS
# ============================================================================

def apply_table_settings(table, settings):
    """
    Apply settings to a table.
    
    Args:
        table: python-docx Table object
        settings: Dict with table properties
    """
    # Alignment
    if 'alignment' in settings:
        alignment = settings['alignment']
        if isinstance(alignment, str):
            alignment_map = {
                'LEFT': WD_TABLE_ALIGNMENT.LEFT,
                'CENTER': WD_TABLE_ALIGNMENT.CENTER,
                'RIGHT': WD_TABLE_ALIGNMENT.RIGHT
            }
            table.alignment = alignment_map.get(alignment.upper(), WD_TABLE_ALIGNMENT.CENTER)
        else:
            table.alignment = alignment
    
    # Autofit
    if 'autofit' in settings:
        table.autofit = settings['autofit']
    
    # Column widths
    if 'column_widths' in settings:
        for i, width in enumerate(settings['column_widths']):
            if i < len(table.columns):
                table.columns[i].width = Cm(width)


def set_row_height(row, height_cm, rule='exact'):
    """
    Set the height of a table row.
    
    Args:
        row: python-docx table row
        height_cm: Height in centimeters
        rule: 'exact' or 'at_least'
    """
    row.height = Cm(height_cm)
    if rule == 'exact':
        row.height_rule = 1  # WD_ROW_HEIGHT_RULE.EXACTLY
    else:
        row.height_rule = 2  # WD_ROW_HEIGHT_RULE.AT_LEAST


def merge_row_cells(row, start_col=0, end_col=None):
    """
    Merge cells in a row.
    
    Args:
        row: python-docx table row
        start_col: Starting column index
        end_col: Ending column index (None = last column)
    """
    cells = row.cells
    if end_col is None:
        end_col = len(cells) - 1
    
    for i in range(end_col, start_col, -1):
        cells[start_col].merge(cells[i])


# ============================================================================
# RICH TEXT HELPERS
# ============================================================================

def create_rich_text_runs(paragraph, segments, base_style=None):
    """
    Create multiple runs with different styles in a paragraph.
    
    Args:
        paragraph: python-docx Paragraph object
        segments: List of dicts with 'text' and optional style overrides
        base_style: Base style config to use
    
    Example:
        segments = [
            {'text': 'Normal text '},
            {'text': 'bold text', 'bold': True},
            {'text': ' and ', 'italic': True},
            {'text': 'colored', 'font_color': 'FF0000'}
        ]
    """
    if base_style is None:
        base_style = get_style_config('document')
    
    for segment in segments:
        text = segment.get('text', '')
        style = base_style.copy()
        
        # Apply segment-specific overrides
        for key, value in segment.items():
            if key != 'text':
                style[key] = value
        
        run = paragraph.add_run(text)
        apply_run_style(run, style)


# ============================================================================
# HEADING FUNCTIONS
# ============================================================================

def create_heading(document, text, level=1, style_overrides=None):
    """
    Create a heading with explicit formatting (no Word style leakage).
    
    Args:
        document: python-docx Document object
        text: Heading text
        level: Heading level (1 or 2)
        style_overrides: Optional style overrides
    
    Returns:
        Created paragraph
    """
    style_name = f'heading{level}'
    style_config = get_style_config(style_name, style_overrides)
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_run_style(run, style_config)
    apply_paragraph_style(paragraph, style_config)
    
    return paragraph
